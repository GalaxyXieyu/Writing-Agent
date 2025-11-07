# -*- coding:utf-8 -*-
import os
import re

import tiktoken
from docx import Document
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
from io import BytesIO
import requests
import json

#计算gpt4使用的tokens
def compute_gpt_tokens(example_string: str):
    encoding = tiktoken.get_encoding("cl100k_base")
    token_integers = encoding.encode(example_string)
    num_tokens = len(token_integers)
    return num_tokens

def check_title_level(txt):
    pattern1 = r'(^\d+、|^[一二三四五六七八九十]+、)'  # 匹配数字或中文数字后面跟随"、"的模式
    matches1 = re.findall(pattern1, txt)    
    if len(matches1) > 0:
        return 1
    pattern2 = r'(^\d+.\d+、)'  # 匹配数字或中文数字后面跟随"、"的模式
    matches2 = re.findall(pattern2, txt)
    if len(matches2) > 0:
        return 2
    pattern3 = r'(^\d+.\d+.\d+、)'  # 匹配数字或中文数字后面跟随"、"的模式
    matches3 = re.findall(pattern3, txt)
    if len(matches3) > 0:
        return 3    
    return None
        
def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    full_text = '\n'.join(full_text)
    len_text_token = len(full_text)*3
    return full_text, len_text_token

def html_to_docx(title, html_content):
    document = Document()
    document.add_heading(title, 0)

    soup = BeautifulSoup(html_content, 'lxml')

    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol']):
        if element.name == 'p':
            document.add_paragraph(element.text)
        elif element.name.startswith('h') and len(element.name) == 2:
            level = int(element.name[1])
            document.add_heading(element.text, level=level)
        elif element.name == 'ul':
            for li in element.find_all('li'):
                document.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                document.add_paragraph(li.text, style='List Number')

    # Handle images
    for img in soup.find_all('img'):
        try:
            response = requests.get(img['src'])
            image_stream = BytesIO(response.content)
            document.add_picture(image_stream, width=Inches(6))
        except:
            print(f"Failed to add image: {img['src']}")

    return document


def extract_template_generate(data):
    try:
        print("需要将数据进行匹配：",data.content)
        match = re.search(r'```json([\s\S]*?)```', data.content)
        if match:
            json_data_string = match.group(1).strip()
            try:
                json_data = json.loads(json_data_string)
                return json_data
            except json.JSONDecodeError as e:
                print(f"解析JSON数据时发生错误: {e}")
        else:
            json_data = json.loads(data.content)
            print("没有找到匹配的JSON数据",json_data)
            return data.content
        print("未找到有效的JSON格式数据")
        content = data.content
        pattern = r'^```json(.*?)```$'
        match = re.search(pattern, content, re.DOTALL)
        if match:
            json_str = match.group(1).strip()
            parsed_data = json.loads(json_str)
            return parsed_data
        else:
            raise ValueError("未找到有效的JSON格式数据")
    except Exception as e:
        return data.content
