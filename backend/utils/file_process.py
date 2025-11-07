import subprocess
from docx import Document
from PyPDF2 import PdfReader
from utils.logger import mylog
import os 
import traceback
from docx import Document
from lxml import etree
from docx.oxml.ns import qn

def html_to_docx(filetitle, html_content):
    # 创建一个新的 DOCX 文档
    doc = Document()
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')  
    # 添加一个标题
    title0 = doc.add_heading(0)
    run = title0.add_run(filetitle)    
    # 设置字体名称，这里设置为微软雅黑
    run.font.name = '微软雅黑'
    # docx库在处理中文时，需要额外设置字体的复杂脚本属性
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')         
    # 使用 lxml 解析 HTML
    tree = etree.HTML(html_content)
    #
    run.font.name = '微软雅黑'
    # docx库在处理中文时，需要额外设置字体的复杂脚本属性
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')      
    # 使用 XPath 选择标签
    tags = tree.xpath('//*[self::p or starts-with(name(), "h")]')
    #print(tags)
    # 提取 HTML 中的文本内容并将其添加到 DOCX 文档中
    for tag in tags:
        if tag.text == None:
            continue
        text = tag.text.strip()
        if tag.tag.startswith('h'):
            levelnum = int(tag.tag[1])
            title = doc.add_heading(level=levelnum)
            run = title.add_run(text)     
            # 设置字体名称，这里设置为微软雅黑
            run.font.name = '微软雅黑'
            # docx库在处理中文时，需要额外设置字体的复杂脚本属性
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')                  
        else:
            doc.add_paragraph("    "+text)#缩进
           
    return doc


def docx_to_pdf(docx_path, pdf_path):
    try:
        subprocess.run(['libreoffice', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path])
        print("pdf文件路径", pdf_path)
        return pdf_path
    except subprocess.CalledProcessError as e:
        print(f"DOCX转PDF失败: {e}")
        return None