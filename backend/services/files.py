import asyncio
import subprocess
from docx import Document
from PyPDF2 import PdfReader
from utils.logger import mylog
import os
import traceback
from ai.agents.planner import make_planner
from ai.llm.llm_factory import LLMFactory
from config import AsyncSessionLocal

# 文本截断上限（与 routes/file.py MAX_FILE_SIZE 对齐）
TRUNCATE_LIMIT = 10000

async def docx_to_pdf(docx_path, pdf_path):
    try:
        # 使用libreoffice将docx文件转换为pdf文件
        subprocess.run(['libreoffice', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path])
        mylog.info(f"PDF文件路径: {pdf_path}")
        return pdf_path
    except subprocess.CalledProcessError as e:
        mylog.error(f"DOCX转PDF失败: {e}")
        return None

async def structure_extract(context):
    try:
        # 说明：此处接入动态模型配置，优先使用默认模型
        # 若无法获取模型，回退本地线程执行（返回基础占位结果）
        loop = asyncio.get_event_loop()
        async with AsyncSessionLocal() as db:
            llm = await LLMFactory.get_default_llm(db)
        if llm is None:
            mylog.warning("structure_extract: 未找到默认模型配置，回退本地执行")
            query = f"请帮我制定该文章{context}的总标题、一级标题、一级标题要求、二级标题、二级标题要求"
            return await loop.run_in_executor(None, lambda: {"titleName": "自动生成标题", "writingRequirement": "", "children": []})
        runnable = make_planner(llm)
        query = f"请帮我制定该文章{context}的总标题、一级标题、一级标题要求、二级标题、二级标题要求"
        # 直接异步调用 llm 可运行体
        return await runnable.ainvoke({"query": query})
    except:
        mylog.info(traceback.format_exc())
        response = {"error": True, "code": 200, "message": "Failed to fetch"}
        return response


async def file_structure_extract(file_path):
    try:
        if not os.path.isfile(file_path):
            mylog.error(f"文件不存在: {file_path}")
            return None

        file_extension = os.path.splitext(file_path)[1].lower()[1:]
        if file_extension == 'pdf':
            # 处理PDF文件
            pdf = PdfReader(file_path)
            num_pages = len(pdf.pages)
            content = [pdf.pages[i].extract_text() for i in range(num_pages)]
            total_characters = sum(len(page) for page in content)
            plan = await structure_extract(content)
            data = {
                'code': 200,
                'type': 'success',
                'pages': num_pages,
                'content': content,
                'fileWords': total_characters,
                'data': plan
            }
            mylog.info(f"PDF文件处理成功: {file_path}, 页数: {num_pages}, 字符总数: {total_characters}")
        elif file_extension == 'docx':
            # 处理DOCX文件
            doc = Document(file_path)
            pdf_path = os.path.splitext(file_path)[0] + '.pdf'
            pdf_path = docx_to_pdf(file_path, pdf_path)
            pdf = PdfReader(pdf_path)
            total_pages = len(pdf.pages)
            full_text = []
            total_words = 0
            for para in doc.paragraphs:
                text = para.text
                full_text.append(text)
                total_words += len(text.split())
            full_text = '\n'.join(full_text)
            plan = await structure_extract(full_text)
            data = {
                'code': 200,
                'type': 'success',
                'pages': total_pages,
                'content': full_text,
                'fileWords': total_words,
                'data': plan
            }
            mylog.info(f"DOCX文件处理成功: {file_path}, 页数: {total_pages}, 单词总数: {total_words}")
        else:
            mylog.error(f"不支持的文件类型: {file_extension}")
            return None

        return data
    except Exception as e:
        error_msg = f"文件解析失败: {str(e)}"
        mylog.error(error_msg)
        mylog.error(traceback.format_exc())
        return None


async def new_file_structure_extract(file_path):
    try:
        if not os.path.isfile(file_path):
            mylog.error(f"文件不存在: {file_path}")
            return None
        file_extension = os.path.splitext(file_path)[1].lower()[1:]
        # 使用线程池处理文件IO操作
        loop = asyncio.get_event_loop()
        if file_extension == 'pdf':
            # 在线程池中执行PDF处理
            def process_pdf():
                pdf = PdfReader(file_path)
                num_pages = len(pdf.pages)
                content = [pdf.pages[i].extract_text() for i in range(num_pages)]
                total_characters = sum(len(page) for page in content)
                return num_pages, content, total_characters
            num_pages, content, total_characters = await loop.run_in_executor(None, process_pdf)
            plan = await structure_extract(content)
            data = {
                'code': 200,
                'type': 'success',
                'pages': num_pages,
                'content': content,
                'fileWords': total_characters,
                'data': plan
            }
            mylog.info(f"PDF文件处理成功: {file_path}, 页数: {num_pages}, 字符总数: {total_characters}")
        elif file_extension == 'docx':
            # 在线程池中执行DOCX处理
            def process_docx():
                full_text = []
                doc = Document(file_path)
                for para in doc.paragraphs:
                    full_text.append(para.text)
                return '\n'.join(full_text)
            
            full_text = await loop.run_in_executor(None, process_docx)
            pdf_path = os.path.splitext(file_path)[0] + '.pdf'
            pdf_path = await docx_to_pdf(file_path, pdf_path)
            pdf = PdfReader(pdf_path)
            total_pages = len(pdf.pages)
            mylog.info(f"DOCX文件处理成功: {full_text}, 页数: {total_pages}")
            # 对 DOCX 文本也做保守截断，避免极端长文占用
            if len(full_text) > TRUNCATE_LIMIT:
                full_text = full_text[:TRUNCATE_LIMIT]
            plan = await structure_extract(full_text)
            data = {
                'code': 200,
                'type': 'success',
                'pages': total_pages,
                'content': full_text,
                'fileWords': len(full_text),
                'data': plan
            }
        elif file_extension in ('txt', 'md'):
            # 读取纯文本/Markdown，并按上限截断
            def read_text():
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        return f.read()
                except Exception:
                    return ''
            full_text = await loop.run_in_executor(None, read_text)
            original_len = len(full_text)
            if original_len > TRUNCATE_LIMIT:
                full_text = full_text[:TRUNCATE_LIMIT]
            mylog.info(f"TEXT/MD 文件处理: {file_path}, 原始字符数: {original_len}, 截断后: {len(full_text)}")
            plan = await structure_extract(full_text)
            data = {
                'code': 200,
                'type': 'success',
                'pages': 1,
                'content': full_text,
                'fileWords': len(full_text),
                'data': plan
            }
        else:
            mylog.error(f"不支持的文件类型: {file_extension}")
            return None

        return data
    except Exception as e:
        error_msg = f"文件解析失败: {str(e)}"
        mylog.error(error_msg)
        mylog.error(traceback.format_exc())
        return None